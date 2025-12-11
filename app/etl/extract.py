from typing import List, Dict
from pathlib import Path
import io
import requests
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
from readability import Document
from lxml import html


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def fetch_url(url: str) -> str:
    r = requests.get(url, timeout=20)
    r.raise_for_status()
    return r.text


def fetch_url_readable(url: str) -> Dict[str, str]:
    headers = {"User-Agent": "Mozilla/5.0 (PersonaMirror)"}
    r = requests.get(url, timeout=30, headers=headers)
    r.raise_for_status()
    doc = Document(r.text)
    summary_html = doc.summary()
    root = html.fromstring(summary_html)
    text = "\n".join([t.strip() for t in root.xpath("//text()") if t.strip()])
    title = doc.title() or url
    return {"source": url, "text": f"{title}\n\n{text}"}


def read_pdf(path: Path) -> str:
    try:
        return pdf_extract_text(str(path))
    except Exception:
        return ""


def read_docx(path: Path) -> str:
    try:
        d = docx.Document(str(path))
        return "\n".join(p.text for p in d.paragraphs)
    except Exception:
        return ""


def extract_from_paths(paths: List[Path]) -> List[Dict[str, str]]:
    items = []
    for p in paths:
        suf = p.suffix.lower()
        if suf in [".txt", ".md"]:
            items.append({"source": str(p), "text": read_text_file(p)})
        elif suf == ".pdf":
            items.append({"source": str(p), "text": read_pdf(p)})
        elif suf == ".docx":
            items.append({"source": str(p), "text": read_docx(p)})
    return items


def extract_from_urls(urls: List[str], readability: bool = True) -> List[Dict[str, str]]:
    items = []
    for u in urls:
        if readability:
            items.append(fetch_url_readable(u))
        else:
            items.append({"source": u, "text": fetch_url(u)})
    return items
